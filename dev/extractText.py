#Extract files

import docx
import re


class TextReader:
    def __init__(self, document):
        self.doc = docx.Document(document)
        self.title = None
        self.body = ""
        self.path = ""

    def compileInfo(self, l, u, title_max, output):
        self.path = output #path for output
        self.title = open(self.path + self.doc.paragraphs[l].runs[0].text, "w+") # initialize doc title
        for i in range(l+2, u):
            if self.doc.paragraphs[i].text != "" and self.doc.paragraphs[i].runs[0].bold != None and not re.match("\w+\/\w+|Defaults|see:", self.doc.paragraphs[i].text): # if title and then next is bold
                print(self.doc.paragraphs[i].runs[0].text)
                print(i)
                try:
                    self.title.write(self.body) # write previous
                except:
                    print("error")
                self.body = "" # clear previous body
                self.title.close() #close file
                try:
                    self.title = open(self.path + self.doc.paragraphs[i].runs[0].text, "w+") # initialize doc title
                except:
                    continue
                continue
            if self.doc.paragraphs[i].text != "" and not self.doc.paragraphs[i].runs[0].bold: #and re.match("^(Variable|(\d+ (points?\/| or)))", self.doc.paragraphs[i].text): # find if text begins with either "variable" or number points 
                self.body = self.body + "\r" + self.doc.paragraphs[i].text
        #final exit
        self.title.write(self.body)
        self.title.close()

        
##doc = docx.Document('BasicSetd.docx')
##for i in range(4497, 4500):
##    if doc.paragraphs[i].text.count(" ") < 4:
##        new = i # beginning of new text
##        title = open("data/output/" + doc.paragraphs[i].text, "w+")
##        title.write("test")
##        title.close()
##        continue
##    
        
    
t = TextReader("BasicSetd.docx")
t.compileInfo(l = 4757, u = 6604, title_max = 4, output = "data/output/")
